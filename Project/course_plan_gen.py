from docx import Document
import pandas as pd
import nltk
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import re
from langchain_community.document_loaders import CSVLoader
from langchain.llms import Ollama
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import spacy
from collections import Counter
from io import StringIO
import tempfile
import streamlit as st


def course_plan_generator():
# Load the Teacher's document using docx
#Extract course outcomes
    co_numbers=[]
    course_outcomes=[]

    st.title("Course Plan Generator")
    st.divider()
    st.header("Please Upload the Syllabus Document here for Processing")
    st.divider()
    uploaded_file = st.file_uploader("Upload a Word document (.docx)", type="docx")
    doc=None
    if uploaded_file is not None:
        doc = Document(uploaded_file)

    for tables in doc.tables:
        for row in tables.rows:
            cells=[cell.text.strip() for cell in row.cells]
            if len(cells)==2:
                co_numbers.append(cells[0])
                course_outcomes.append(cells[1])

    # Directly create DataFrame from extracted table data if column headings are included
    df = pd.DataFrame([co_numbers, course_outcomes]).transpose()

    # Rename columns only if needed
    df.columns = df.iloc[0]  # Set the first row as the header
    df = df[1:].reset_index(drop=True) 

    # Extracting the Units - Topics - Teaching Hours from syllabus
    units = []
    current_unit = None
    current_content = []

    # Parse through the document paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()

        # Check for "Unit" and start a new unit
        if text.startswith("Unit"):
            # Save the previous unit and its content
            if current_unit:
                units.append((current_unit, " ".join(current_content)))
            
            # Start a new unit
            current_unit = text
            current_content = []
        elif current_unit:
            # Check if the paragraph contains Lab Exercises or Reading sections
            if text.startswith("Lab Exercise") or text.startswith("Essential Reading") or text.startswith("Recommended Reading"):
                continue
            # Accumulate content for the current unit
            current_content.append(text)

    # Append the last unit
    if current_unit:
        units.append((current_unit, " ".join(current_content)))

    # Create a DataFrame with columns "Unit", "Contents"
    df_units = pd.DataFrame(units, columns=["Unit", "Contents"])

    # Extract teaching hours using the specific pattern "Teaching Hours: X"
    def extract_hours(contents):
        match = re.search(r"Teaching Hours:\s*(\d+)", contents)
        return int(match.group(1)) if match else None

    # Extract content before "Teaching Hours"
    def extract_content_before_hours(contents):
        if "Teaching Hours" in contents:
            return contents.split("Teaching Hours")[0].strip()
        return contents.strip()

    # Apply content splitting and teaching hours extraction
    df_units['Teaching Hours'] = df_units['Contents'].apply(extract_hours)
    df_units['Contents'] = df_units['Contents'].apply(extract_content_before_hours)

    # Extract Topic from the unit by assuming it's the part of the string after "Unit X:"
    def extract_topic(unit):
        # Match unit topic patterns with different possible delimiters
        match = re.search(r"Unit\s*\d+\s*[:\t\s](.+)", unit)
        return match.group(1).strip() if match else ""  # Return empty string if no match


    # Apply topic extraction
    df_units['Topic'] = df_units['Unit'].apply(extract_topic)

    # Clean up "Unit" column to only contain the unit number (e.g., "Unit 1")
    df_units['Unit'] = df_units['Unit'].apply(lambda x: re.match(r"Unit\s*\d+", x).group())

    # Reorder columns for better readability
    df_units = df_units[['Unit', 'Topic', 'Contents', 'Teaching Hours']]

    ### Matching the units with their respective course outcomes based on the similarity score
    from sklearn.feature_extraction.text import TfidfVectorizer

    # Preprocess text (lowercase for consistency)
    df['Course Outcomes'] = df['Course Outcomes'].str.lower()
    df_units['Topic'] = df_units['Topic'].str.lower()

    # Combine text data for vectorization
    all_text = pd.concat([df['Course Outcomes'], df_units['Topic']])

    # Vectorize using TF-IDF
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(all_text)

    # Separate vectors for Course Outcomes and Topics
    course_outcome_vectors = tfidf_matrix[:len(df)]
    topic_vectors = tfidf_matrix[len(df):]

    # Compute cosine similarity
    similarity_matrix = cosine_similarity(topic_vectors, course_outcome_vectors)

    # Find the most similar course outcome for each unit
    best_matches = []
    for i, topic in enumerate(df_units['Topic']):
        # Get the index of the most similar course outcome
        best_match_index = similarity_matrix[i].argmax()
        best_match_score = similarity_matrix[i][best_match_index]
        
        # Append the best match and its score
        best_matches.append({
            'Matched Course Outcome': df.iloc[best_match_index]['Course Outcomes'],
            'Similarity Score': best_match_score
        })

    # Convert matches to a DataFrame
    matches_df = pd.DataFrame(best_matches)

    # Add the matched course outcome and score to df_units
    df_units['Course Outcomes'] = matches_df['Matched Course Outcome']
    df_units['Similarity Score'] = matches_df['Similarity Score']

    # Display the updated DataFrame
    print(df_units)

    # Extract Verbs from the Course Outcomes
    course_outcomes=[]
    for i in range(len(df)):
        data=df_units['Course Outcomes'].iloc[i]
        course_outcomes.append(data)

    #Verbs
    verbs=['VB','VBP','VBD','VBG','VBN']
    course_verbs = []
    for i in range(len(course_outcomes)):
        review = course_outcomes[i]
        review = review.split()
        review = nltk.pos_tag(review)
        filtered_verbs = [word for word, tag in review if tag in verbs]
        course_verbs.append(filtered_verbs)

    # Compare lengths of df_units and course_verbs
    if len(course_verbs) < len(df_units):
    # Extend course_verbs to match df_units length
        course_verbs.extend([[]] * (len(df_units) - len(course_verbs)))
    elif len(course_verbs) > len(df_units):
    # Truncate course_verbs to match df_units length
        course_verbs = course_verbs[:len(df_units)]

# Assign adjusted course_verbs to the 'Verbs' column
    df_units['Verbs'] = course_verbs

    # Load the defined verb-assesment dataset (grouped)
    df_assessments=pd.read_excel("Verbs-Assesments Grouped.xlsx")
    # Convert the 'verbs' column to lists of individual words, ignoring NaNs
    flattened_assessment_verbs = []
    for item in df_assessments['Verbs'].dropna():
        flattened_assessment_verbs.extend([verb.strip() for verb in item.split(',')])

    nlp = spacy.load("en_core_web_md")

    # Function to get the SpaCy vector for a word
    def get_word_vector(word):
        doc = nlp(word)
        if doc.has_vector:
            return doc.vector
        else:
            return np.zeros(nlp.vocab.vectors_length)  # Return zero vector if word not in vocabulary

    # Expand 'Verbs' column to match a flattened list
    expanded_rows = []
    for index, row in df_assessments.iterrows():
        verbs = row['Verbs']
        if pd.isna(verbs):  # Skip rows with NaN in 'Verbs'
            continue
        verb_list = [verb.strip() for verb in verbs.split(',')]  # Split multiple verbs into a list
        for verb in verb_list:
            expanded_rows.append({'Verbs': verb, 'Assessments': row['Assessments']})

    # Create a new expanded DataFrame
    expanded_df_assessments = pd.DataFrame(expanded_rows)

    #  Match and extract the assesments for the verbs extracted from the syllabus CO
    flattened_course_verbs = [verb for sublist in course_verbs for verb in sublist]
    # Flattened assessment verbs
    flattened_assessment_verbs = expanded_df_assessments['Verbs'].tolist()

    # Convert course verbs and assessment verbs to vectors
    course_verb_vectors = np.array([get_word_vector(verb) for verb in flattened_course_verbs])
    assessment_verb_vectors = np.array([get_word_vector(verb) for verb in flattened_assessment_verbs])

    # Similarity threshold
    threshold = 0.5

    # Dictionary to store matching assessments
    matching_assessments = {}

    # Function to normalize assessments (you can extend this for more complex cases)
    def normalize_assessment(assessment):
        # Strip leading/trailing spaces and convert to lowercase
        return assessment.strip().lower()

    for i, course_verb in enumerate(flattened_course_verbs):
        matching_assessments[course_verb] = set()  # Use a set to avoid duplicate assessments

        # Calculate cosine similarity between the course verb and each assessment verb
        sim_scores = cosine_similarity([course_verb_vectors[i]], assessment_verb_vectors).flatten()

        # Find assessment verbs with similarity scores above the threshold
        for j, score in enumerate(sim_scores):
            if score > threshold:
                if 0 <= j < len(expanded_df_assessments):  # Ensure index is valid
                    normalized_assessment = normalize_assessment(expanded_df_assessments.iloc[j]['Assessments'])
                    matching_assessments[course_verb].add(normalized_assessment)  # Store normalized assessment

    # Print results
    # Create an empty DataFrame with specified columns
    df_matchass = pd.DataFrame(columns=["Verbs", "Assessments"])

    # Iterate through matching assessments
    for verb, assessments in matching_assessments.items():
        print(f"Verb: {verb}")
        if assessments:
            print("Matching Assessments:")
            # Sort and remove duplicates (optional, for better display)
            unique_assessments = sorted(set(assessments))
            print("\n".join(f"- {assessment}" for assessment in unique_assessments))
            
            # Create a new row as a DataFrame
            new_row = pd.DataFrame({
                "Verbs": [verb],
                "Assessments": [", ".join(unique_assessments)]
            })
            
            # Concatenate the new row to the existing DataFrame
            df_matchass = pd.concat([df_matchass, new_row], ignore_index=True)
        else:
            print("No matching assessments found.")
        print("\n")

    assessments = []

    # Iterate over each row in df_units
    for index, row in df_units.iterrows():
        # Initialize a list to store matched assessments for each verb list
        matched_assessments = []
        
        # Iterate over the list of verbs in the "Verbs" column of df_units
        for verb in row['Verbs']:
            # Escape special characters in the verb for regex matching
            escaped_verb = re.escape(verb)
            
            # Check if the verb exists in df_matchass
            matching_rows = df_matchass[df_matchass['Verbs'].str.contains(escaped_verb, case=False, na=False)]
            
            # If there are matches, collect the corresponding assessments
            if not matching_rows.empty:
                matched_assessments.extend(matching_rows['Assessments'].tolist())
        
        # Remove duplicate assessments if any and add to the list
        matched_assessments = list(set(matched_assessments))
        
        # Append the matched assessments to the assessments list
        assessments.append(matched_assessments)

    # Add the new "Assessment" column to df_units
    df_units['Assessments'] = assessments

    # Convert the list of characters in the "Verbs" column back to strings
    df_units['Verbs'] = df_units['Verbs'].apply(lambda x: ''.join(x) if isinstance(x, list) else x)

    # Getting only unique assessments - Filtering

    def clean_assessments(assessments):
        if isinstance(assessments, list):  # Check if it's already a list
            assessments = " ".join(assessments)  # Convert the list to a single string
        assessments = assessments.strip("[]").replace("'", "").replace('"', "")
        return [x.strip() for x in assessments.split("•") if x.strip()]

    df_units["Assessments_cleaned"] = df_units["Assessments"].apply(clean_assessments)


    all_assessments = [item for sublist in df_units["Assessments_cleaned"] for item in sublist]
    assessment_counts = Counter(all_assessments)

    # Step 4: Apply the filtering logic iteratively
    def filter_assessments(row, assessment_counts, threshold=3):
        unique_assessments = []
        
        for assessment in row:
            # If the assessment appears more than once and the unit has more than threshold assessments, remove it
            if assessment_counts[assessment] > 1 and len(row) > threshold:
                continue  # Remove this assessment from the unit
            # If the assessment appears more than once and the unit has less than threshold assessments, keep it
            elif assessment_counts[assessment] > 1 and len(row) < threshold:
                unique_assessments.append(assessment)
            # Otherwise, always add the assessment
            else:
                unique_assessments.append(assessment)
        
        return unique_assessments

    # Step 5: Apply the filtering function to each row (unit)
    df_units["Filtered Assessments"] = df_units["Assessments_cleaned"].apply(
        lambda row: filter_assessments(row, assessment_counts, threshold=3)
    )


    # Step 4: Apply the filtering logic iteratively
    def filter_assessments(row, assessment_counts, threshold=3):
        # Create a list to hold filtered assessments
        filtered_assessments = []
        # Iterate over the assessments of each unit
        for assessment in row:
            # If the assessment appears more than once and the unit has more than threshold assessments, remove it
            if assessment_counts[assessment] > 1 and len(row) > threshold:
                continue  # Remove this assessment from the unit
            # If the assessment appears more than once and the unit has less than threshold assessments, keep it
            elif assessment_counts[assessment] > 1 and len(row) < threshold:
                filtered_assessments.append(assessment)
            # Otherwise, always add the assessment
            else:
                filtered_assessments.append(assessment)
        
        return filtered_assessments

    # Step 5: Apply the filtering function to each row (unit)
    df_units["Filtered Assessments"] = df_units["Assessments_cleaned"].apply(
        lambda row: filter_assessments(row, assessment_counts, threshold=3)
    )

    # Step 6: Ensure no unit is left empty but limit to top 3 assessments if necessary
    def recheck_and_limit_empty_units(df, limit=3):
        for index, row in df.iterrows():
            # If a unit's filtered assessments are empty, we restore the top 3 from its original assessments
            if not row["Filtered Assessments"]:
                original_assessments = row["Assessments_cleaned"]
                # Ensure that the filtered assessments are limited to the top 3
                df.at[index, "Filtered Assessments"] = original_assessments[:limit]
        return df

    # Step 7: Apply the recheck to limit to top 3 assessments
    df_units = recheck_and_limit_empty_units(df_units)
    df_units.drop(['Assessments','Assessments_cleaned'],axis=1,inplace=True)



    def create_temp_csv_for_langchain(df):
        """
        Creates a temporary CSV file from a DataFrame and returns its path.

        Parameters:
            df (pd.DataFrame): The DataFrame to convert.

        Returns:
            str: The file path of the temporary CSV file.
        """
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="w")  # Create a temp file
        df.to_csv(temp_file.name, index=False)  # Save the DataFrame as a CSV to the temp file
        temp_file.close()  # Close the file to allow external access
        return temp_file.name

    #Loading using Langchain Loader

    temp_csv_path = create_temp_csv_for_langchain(df_units)

    loader = CSVLoader(file_path=temp_csv_path, csv_args={
        'delimiter': ',',
        'quotechar': '"',
        
    })

    documents = loader.load()



    # Step 2: Extract structured data from loaded documents
    data = []
    for doc in documents:
        content_lines = doc.page_content.strip().split('\n')
    
        if len(content_lines) >= 7:  # Adjusted to match the actual structure
            entry = {
                "Unit": content_lines[0].replace("Unit:", "").strip(),
                "Topic": content_lines[1].replace("Topic:", "").strip(),
                "Contents": content_lines[2].replace("Contents:", "").strip(),
                "Teaching Hours": content_lines[3].replace("Teaching Hours:", "").strip(),
                "Course Outcomes": content_lines[4].replace("Course Outcomes:", "").strip(),
                "Similarity Score": content_lines[5].replace("Similarity Score:", "").strip(),
                "Verbs": content_lines[6].replace("Verbs:", "").strip(),
                "Filtered Assessments": content_lines[7].replace("Filtered Assessments:", "").strip(),
            }
            data.append(entry)

    if data:
    # Extract unique units
        unique_units = sorted(set(item["Unit"] for item in data if "Unit" in item))

        # Display title for all units
        st.write("## Details for All Units")

        # Collect data for all units
        all_units_data = []

        # Loop through each unit and collect its data
        for unit in unique_units:
            st.write(f"### Unit: {unit}")

            # Filter data for the current unit
            unit_data = [item for item in data if item["Unit"] == unit]
            
            if unit_data:
                # Convert to a DataFrame for better presentation
                display_data = pd.DataFrame(unit_data)

                # Extract required fields for further processing
                topics = ", ".join(item["Topic"] for item in unit_data)
                contents = ", ".join(item["Contents"] for item in unit_data)
                teaching_hours = ", ".join(item["Teaching Hours"] for item in unit_data)
                assessments = ", ".join(item["Filtered Assessments"] for item in unit_data)

                # Collect this unit's data into all_units_data
                all_units_data.append({
                    "unit": unit,
                    "contents": contents,
                    "teaching_hours": teaching_hours,
                    "assessments": assessments
                })

                # Display extracted details clearly
                st.write("#### Topics")
                st.write(topics)

                st.write("#### Contents")
                st.write(contents)

                st.write("#### Teaching Hours")
                st.write(teaching_hours)

                st.write("#### Assessments")
                st.write(assessments)

                st.write("---")  # Add a separator between units
                
            else:
                st.error(f"No data found for unit: {unit}")

        # Step 2: Generate the course plan for all units combined
        if st.button("Generate Course Plan for All Units"):
            st.write("Generating Course Plan. Please wait...")

            # Define the prompt template
            prompt_template = PromptTemplate(
                input_variables=["all_units_data"],
                template=(  """
   You are an AI educational planner tasked with generating a comprehensive and structured course plan for multiple units. Below is the combined information for all units:

    {all_units_data}

Task:
1. Each session will be treated as 1 individual hour.
2. For each session, specify:
   - Session Number (e.g., Session 1, Session 2, etc.)
   - Content to be covered during the session, ensuring it aligns with the total unit contents. Break the content logically across the sessions to maintain a smooth flow.
   - Duration: 1 hour.
   - Key Topics/Concepts to be covered in the session.

3. The course plan should aim for an even distribution of teaching hours across all sessions. If necessary, prioritize core concepts earlier in the plan and ensure the pacing is manageable for students.

Example Output:
- Session 1:
   - Duration: 1 Hour
   - Content: Introduction to word embeddings and vector representations
   - Topics: Basic definition, importance, and key concepts of word embeddings

- Session 2:
   - Duration: 1 Hour
   - Content: Applications of word embeddings in NLP
   - Topics: Search engines, semantic similarity, and sentiment analysis

Deliverable:
Provide a detailed course plan with an appropriate number of 1-hour sessions for all units, each covering part of the unit's content. Ensure alignment between the session duration, total teaching hours, and content while maintaining logical and manageable pacing for students.
                            """
                    )
                )

            # Initialize LLM and chain
            llama = Ollama(model="llama3.2")
            llm_chain = LLMChain(llm=llama, prompt=prompt_template)

            # Run the prompt with all units' data
            response = llm_chain.run({
                "all_units_data": all_units_data
            })

            # Display the generated course plan for all units
            st.write("### Generated Course Plan for All Units")
            st.text(response)
            
            # Option to download the results
            st.download_button(
                label="Download Course Plan for All Units",
                data=response,
                file_name="all_units_course_plan.txt",
                mime="text/plain"
            )
        else:
            st.error("No data available to display. Please load valid structured data.")
