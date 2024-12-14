from docx import Document
import pandas as pd
import nltk
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import re
from langchain_community.document_loaders import CSVLoader
from langchain.llms import Ollama
from langchain.llms import LLMChain
from langchain.prompts import PromptTemplate
import spacy
from collections import Counter
from io import StringIO
import tempfile

# Load the Teacher's document using docx
#Extract course outcomes
co_numbers=[]
course_outcomes=[]

doc=Document('Syllabus.docx')

for tables in doc.tables:
    for row in tables.rows:
        cells=[cell.text.strip() for cell in row.cells]
        if len(cells)==2:
            co_numbers.append(cells[0])
            course_outcomes.append(cells[1])

# Directly create DataFrame from extracted table data if column headings are included
df = pd.DataFrame([co_numbers, course_outcomes]).transpose()

# Rename columns only if needed
df.columns = df.iloc[0]  # Set the first row as the header
df = df[1:].reset_index(drop=True) 

# Extracting the Units - Topics - Teaching Hours from syllabus
units = []
current_unit = None
current_content = []

# Parse through the document paragraphs
for para in doc.paragraphs:
    text = para.text.strip()

    # Check for "Unit" and start a new unit
    if text.startswith("Unit"):
        # Save the previous unit and its content
        if current_unit:
            units.append((current_unit, " ".join(current_content)))
        
        # Start a new unit
        current_unit = text
        current_content = []
    elif current_unit:
        # Check if the paragraph contains Lab Exercises or Reading sections
        if text.startswith("Lab Exercise") or text.startswith("Essential Reading") or text.startswith("Recommended Reading"):
            continue
        # Accumulate content for the current unit
        current_content.append(text)

# Append the last unit
if current_unit:
    units.append((current_unit, " ".join(current_content)))

# Create a DataFrame with columns "Unit", "Contents"
df_units = pd.DataFrame(units, columns=["Unit", "Contents"])

# Extract teaching hours using the specific pattern "Teaching Hours: X"
def extract_hours(contents):
    match = re.search(r"Teaching Hours:\s*(\d+)", contents)
    return int(match.group(1)) if match else None

# Extract content before "Teaching Hours"
def extract_content_before_hours(contents):
    if "Teaching Hours" in contents:
        return contents.split("Teaching Hours")[0].strip()
    return contents.strip()

# Apply content splitting and teaching hours extraction
df_units['Teaching Hours'] = df_units['Contents'].apply(extract_hours)
df_units['Contents'] = df_units['Contents'].apply(extract_content_before_hours)

# Extract Topic from the unit by assuming it's the part of the string after "Unit X:"
def extract_topic(unit):
    # Match unit topic patterns with different possible delimiters
    match = re.search(r"Unit\s*\d+\s*[:\t\s](.+)", unit)
    return match.group(1).strip() if match else ""  # Return empty string if no match


# Apply topic extraction
df_units['Topic'] = df_units['Unit'].apply(extract_topic)

# Clean up "Unit" column to only contain the unit number (e.g., "Unit 1")
df_units['Unit'] = df_units['Unit'].apply(lambda x: re.match(r"Unit\s*\d+", x).group())

# Reorder columns for better readability
df_units = df_units[['Unit', 'Topic', 'Contents', 'Teaching Hours']]

### Matching the units with their respective course outcomes based on the similarity score
from sklearn.feature_extraction.text import TfidfVectorizer

# Preprocess text (lowercase for consistency)
df['Course Outcomes'] = df['Course Outcomes'].str.lower()
df_units['Topic'] = df_units['Topic'].str.lower()

# Combine text data for vectorization
all_text = pd.concat([df['Course Outcomes'], df_units['Topic']])

# Vectorize using TF-IDF
vectorizer = TfidfVectorizer()
tfidf_matrix = vectorizer.fit_transform(all_text)

# Separate vectors for Course Outcomes and Topics
course_outcome_vectors = tfidf_matrix[:len(df)]
topic_vectors = tfidf_matrix[len(df):]

# Compute cosine similarity
similarity_matrix = cosine_similarity(topic_vectors, course_outcome_vectors)

# Find the most similar course outcome for each unit
best_matches = []
for i, topic in enumerate(df_units['Topic']):
    # Get the index of the most similar course outcome
    best_match_index = similarity_matrix[i].argmax()
    best_match_score = similarity_matrix[i][best_match_index]
    
    # Append the best match and its score
    best_matches.append({
        'Matched Course Outcome': df.iloc[best_match_index]['Course Outcomes'],
        'Similarity Score': best_match_score
    })

# Convert matches to a DataFrame
matches_df = pd.DataFrame(best_matches)

# Add the matched course outcome and score to df_units
df_units['Course Outcomes'] = matches_df['Matched Course Outcome']
df_units['Similarity Score'] = matches_df['Similarity Score']

# Display the updated DataFrame
print(df_units)

# Extract Verbs from the Course Outcomes
course_outcomes=[]
for i in range(len(df)+1):
    data=df_units['Course Outcomes'].iloc[i]
    course_outcomes.append(data)

#Verbs
verbs=['VB','VBP','VBD','VBG','VBN']
course_verbs = []
for i in range(len(course_outcomes)):
    review = course_outcomes[i]
    review = review.split()
    review = nltk.pos_tag(review)
    filtered_verbs = [word for word, tag in review if tag in verbs]
    course_verbs.append(filtered_verbs)

# Assign the collected verbs list to the DataFrame column
df_units['Verbs'] = course_verbs

# Load the defined verb-assesment dataset (grouped)
df_assessments=pd.read_excel("Verbs-Assesments Grouped.xlsx")
# Convert the 'verbs' column to lists of individual words, ignoring NaNs
flattened_assessment_verbs = []
for item in df_assessments['Verbs'].dropna():
    flattened_assessment_verbs.extend([verb.strip() for verb in item.split(',')])

nlp = spacy.load("en_core_web_md")

# Function to get the SpaCy vector for a word
def get_word_vector(word):
    doc = nlp(word)
    if doc.has_vector:
        return doc.vector
    else:
        return np.zeros(nlp.vocab.vectors_length)  # Return zero vector if word not in vocabulary

# Expand 'Verbs' column to match a flattened list
expanded_rows = []
for index, row in df_assessments.iterrows():
    verbs = row['Verbs']
    if pd.isna(verbs):  # Skip rows with NaN in 'Verbs'
        continue
    verb_list = [verb.strip() for verb in verbs.split(',')]  # Split multiple verbs into a list
    for verb in verb_list:
        expanded_rows.append({'Verbs': verb, 'Assessments': row['Assessments']})

# Create a new expanded DataFrame
expanded_df_assessments = pd.DataFrame(expanded_rows)

#  Match and extract the assesments for the verbs extracted from the syllabus CO
flattened_course_verbs = [verb for sublist in course_verbs for verb in sublist]
# Flattened assessment verbs
flattened_assessment_verbs = expanded_df_assessments['Verbs'].tolist()

# Convert course verbs and assessment verbs to vectors
course_verb_vectors = np.array([get_word_vector(verb) for verb in flattened_course_verbs])
assessment_verb_vectors = np.array([get_word_vector(verb) for verb in flattened_assessment_verbs])

# Similarity threshold
threshold = 0.5

# Dictionary to store matching assessments
matching_assessments = {}

# Function to normalize assessments (you can extend this for more complex cases)
def normalize_assessment(assessment):
    # Strip leading/trailing spaces and convert to lowercase
    return assessment.strip().lower()

for i, course_verb in enumerate(flattened_course_verbs):
    matching_assessments[course_verb] = set()  # Use a set to avoid duplicate assessments

    # Calculate cosine similarity between the course verb and each assessment verb
    sim_scores = cosine_similarity([course_verb_vectors[i]], assessment_verb_vectors).flatten()

    # Find assessment verbs with similarity scores above the threshold
    for j, score in enumerate(sim_scores):
        if score > threshold:
            if 0 <= j < len(expanded_df_assessments):  # Ensure index is valid
                normalized_assessment = normalize_assessment(expanded_df_assessments.iloc[j]['Assessments'])
                matching_assessments[course_verb].add(normalized_assessment)  # Store normalized assessment

# Print results
# Create an empty DataFrame with specified columns
df_matchass = pd.DataFrame(columns=["Verbs", "Assessments"])

# Iterate through matching assessments
for verb, assessments in matching_assessments.items():
    print(f"Verb: {verb}")
    if assessments:
        print("Matching Assessments:")
        # Sort and remove duplicates (optional, for better display)
        unique_assessments = sorted(set(assessments))
        print("\n".join(f"- {assessment}" for assessment in unique_assessments))
        
        # Create a new row as a DataFrame
        new_row = pd.DataFrame({
            "Verbs": [verb],
            "Assessments": [", ".join(unique_assessments)]
        })
        
        # Concatenate the new row to the existing DataFrame
        df_matchass = pd.concat([df_matchass, new_row], ignore_index=True)
    else:
        print("No matching assessments found.")
    print("\n")

assessments = []

# Iterate over each row in df_units
for index, row in df_units.iterrows():
    # Initialize a list to store matched assessments for each verb list
    matched_assessments = []
    
    # Iterate over the list of verbs in the "Verbs" column of df_units
    for verb in row['Verbs']:
        # Escape special characters in the verb for regex matching
        escaped_verb = re.escape(verb)
        
        # Check if the verb exists in df_matchass
        matching_rows = df_matchass[df_matchass['Verbs'].str.contains(escaped_verb, case=False, na=False)]
        
        # If there are matches, collect the corresponding assessments
        if not matching_rows.empty:
            matched_assessments.extend(matching_rows['Assessments'].tolist())
    
    # Remove duplicate assessments if any and add to the list
    matched_assessments = list(set(matched_assessments))
    
    # Append the matched assessments to the assessments list
    assessments.append(matched_assessments)

# Add the new "Assessment" column to df_units
df_units['Assessments'] = assessments

# Convert the list of characters in the "Verbs" column back to strings
df_units['Verbs'] = df_units['Verbs'].apply(lambda x: ''.join(x) if isinstance(x, list) else x)

# Getting only unique assessments - Filtering

def clean_assessments(assessments):
    if isinstance(assessments, list):  # Check if it's already a list
        assessments = " ".join(assessments)  # Convert the list to a single string
    assessments = assessments.strip("[]").replace("'", "").replace('"', "")
    return [x.strip() for x in assessments.split("•") if x.strip()]

df_units["Assessments_cleaned"] = df_units["Assessments"].apply(clean_assessments)


all_assessments = [item for sublist in df_units["Assessments_cleaned"] for item in sublist]
assessment_counts = Counter(all_assessments)

# Step 4: Apply the filtering logic iteratively
def filter_assessments(row, assessment_counts, threshold=3):
    unique_assessments = []
    
    for assessment in row:
        # If the assessment appears more than once and the unit has more than threshold assessments, remove it
        if assessment_counts[assessment] > 1 and len(row) > threshold:
            continue  # Remove this assessment from the unit
        # If the assessment appears more than once and the unit has less than threshold assessments, keep it
        elif assessment_counts[assessment] > 1 and len(row) < threshold:
            unique_assessments.append(assessment)
        # Otherwise, always add the assessment
        else:
            unique_assessments.append(assessment)
    
    return unique_assessments

# Step 5: Apply the filtering function to each row (unit)
df_units["Filtered Assessments"] = df_units["Assessments_cleaned"].apply(
    lambda row: filter_assessments(row, assessment_counts, threshold=3)
)


# Step 4: Apply the filtering logic iteratively
def filter_assessments(row, assessment_counts, threshold=3):
    # Create a list to hold filtered assessments
    filtered_assessments = []
    # Iterate over the assessments of each unit
    for assessment in row:
        # If the assessment appears more than once and the unit has more than threshold assessments, remove it
        if assessment_counts[assessment] > 1 and len(row) > threshold:
            continue  # Remove this assessment from the unit
        # If the assessment appears more than once and the unit has less than threshold assessments, keep it
        elif assessment_counts[assessment] > 1 and len(row) < threshold:
            filtered_assessments.append(assessment)
        # Otherwise, always add the assessment
        else:
            filtered_assessments.append(assessment)
    
    return filtered_assessments

# Step 5: Apply the filtering function to each row (unit)
df_units["Filtered Assessments"] = df_units["Assessments_cleaned"].apply(
    lambda row: filter_assessments(row, assessment_counts, threshold=3)
)

# Step 6: Ensure no unit is left empty but limit to top 3 assessments if necessary
def recheck_and_limit_empty_units(df, limit=3):
    for index, row in df.iterrows():
        # If a unit's filtered assessments are empty, we restore the top 3 from its original assessments
        if not row["Filtered Assessments"]:
            original_assessments = row["Assessments_cleaned"]
            # Ensure that the filtered assessments are limited to the top 3
            df.at[index, "Filtered Assessments"] = original_assessments[:limit]
    return df

# Step 7: Apply the recheck to limit to top 3 assessments
df_units = recheck_and_limit_empty_units(df_units)



def create_temp_csv_for_langchain(df):
    """
    Creates a temporary CSV file from a DataFrame and returns its path.

    Parameters:
        df (pd.DataFrame): The DataFrame to convert.

    Returns:
        str: The file path of the temporary CSV file.
    """
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="w")  # Create a temp file
    df.to_csv(temp_file.name, index=False)  # Save the DataFrame as a CSV to the temp file
    temp_file.close()  # Close the file to allow external access
    return temp_file.name

#Loading using Langchain Loader

temp_csv_path = create_temp_csv_for_langchain(df)

loader = CSVLoader(file_path=temp_csv_path, csv_args={
    'delimiter': ',',
    'quotechar': '"',
    'fieldnames': ['','Unit','Topic','Contents','Teaching Hours','Course Outcomes','Similarity Score','Verbs','Assessments']
})

documents = loader.load()



# Step 2: Extract structured data from loaded documents
data = []
for doc in documents:
    # Extract content from each document
    content = doc.page_content.strip()
    
    # Split the content by newline characters to separate columns
    content_lines = content.split('\n')
    
    # Ensure there are enough parts in the split content
    if len(content_lines) >= 9:
        entry = {
            "Unit": content_lines[0].strip(),          # First line: Unit
            "Topic": content_lines[2].strip(),         # Second line: Topic
            "Contents": content_lines[3].strip(),      # Third line: Contents
            "Assessments": set(content_lines[8].strip().split(';'))  # Ninth line: Assessments (handling repeated assessments)
        }
        # Join the assessments back into a single string (if they were split into a set)
        entry["Assessments"] = ', '.join(entry["Assessments"])
        
        data.append(entry)



# Step 3: Define a prompt template
prompt_template = PromptTemplate(
    input_variables=["unit", "contents", "assessments"],
    template=(
        """
You are an AI specializing in educational assessments, designed to work through tasks step by step to generate targeted and comprehensive assessments. Below is information about a unit:
Unit: {unit}
Contents: {contents}
Vague Assessments: {assessments}

Task:

Identify Assessment Types: Analyze the provided vague assessments and identify a broad range of specific assessment types, such as 'essay,' 'MCQ,' 'fill in the blank,' 'case study,' 'diagram labeling,' 'problem-solving,' 'coding exercises,' or others.
Link to Content: Extract the key topics and concepts from the unit's content that align with the identified assessment types.
Generate Tailored Assessments: For each assessment type, create detailed and actionable assessments based on the unit's content. Examples include:
Essay: If the content includes 'word embeddings,' the assessment could be: "Write a detailed essay explaining the concept of word embeddings and their role in natural language processing."
Fill in the Blank: For the topic 'vector representations,' the assessment could be: "_____ is the process of representing words as vectors in a high-dimensional space."
Case Study: If the content covers 'applications of word embeddings,' the assessment could be: "Analyze a case study where word embeddings were used to improve search engine results."
Diagram Labeling: If the content involves 'neural network structures,' the assessment could be: "Label the components of the diagram showing the architecture of a word embedding model."
Problem-Solving: For a topic on 'cosine similarity,' the assessment could be: "Calculate the cosine similarity between the following two word vectors."
Deliverable: Provide diverse, specific, and measurable assessments tailored to the unit, ensuring alignment with the content and coverage of a wide range of assessment types.
"""
    )
)


# Step 4: Initialize the LLM and create a chain
llama = Ollama(model="llama3.2")
llm_chain = LLMChain(llm=llama, prompt=prompt_template)

# Step 5: Generate specific assessments
specific_assessments = []
for entry in data:
    response = llm_chain.run({ 
        "unit": entry["Unit"],
        "contents": entry["Contents"],
        "assessments": entry["Assessments"]
    })
    specific_assessments.append({"Unit": entry["Unit"], "Specific Assessments": response})

# Step 6: Save the results to a CSV file
df = pd.DataFrame(specific_assessments)
output_file = "specific_assessments.csv"
df.to_csv(output_file, index=False)

print(f"Specific assessments have been saved to {output_file}")




